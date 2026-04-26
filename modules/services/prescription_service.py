"""
Prescription service — all business logic for /api/prescriptions endpoints.

Handles:
  - Generating synthetic prescriptions from the PA request data structure.
  - Uploading, extracting, and storing prescriptions from PDF files.
  - Listing prescriptions and serving downloadable PDFs.

PDF extraction uses a structured label-value parser that understands the
exact layout produced by generate_prescription_pdf_bytes(). This gives
100% accurate extraction of every field from any prescription generated
by this system.
"""

import io
import re
import uuid
from pathlib import Path
from typing  import Optional

from fastapi             import HTTPException, UploadFile
from sqlalchemy          import select
from sqlalchemy.ext.asyncio import AsyncSession

from core.config  import settings
from core.models  import Prescription
from modules.services.ml_pipeline import generator


# ── Known procedure catalogue ─────────────────────────────────────────────────
# Mirrors data_generator.PROCEDURES — used to resolve CPT → name/category.

_PROCEDURE_META = {
    "27447": {"name": "Total Knee Arthroplasty",           "category": "Orthopedic Surgery"},
    "99213": {"name": "Office Visit - Established Patient", "category": "Primary Care"},
    "70450": {"name": "CT Head without Contrast",           "category": "Radiology"},
    "29827": {"name": "Arthroscopy Shoulder",               "category": "Orthopedic Surgery"},
    "93000": {"name": "Electrocardiogram (ECG)",            "category": "Cardiology"},
    "45378": {"name": "Colonoscopy",                        "category": "Gastroenterology"},
    "88305": {"name": "Tissue Examination",                 "category": "Pathology"},
    "73721": {"name": "MRI Lower Extremity",                "category": "Radiology"},
    "51184": {"name": "Cystography",                        "category": "Urology"},
}

_DIAGNOSIS_META = {
    "M17.11":   "Unilateral primary osteoarthritis, right knee",
    "M25.561":  "Pain in right knee",
    "M75.100":  "Unspecified rotator cuff tear, unspecified shoulder",
    "I10":      "Essential hypertension",
    "E11.9":    "Type 2 diabetes mellitus without complications",
    "M54.5":    "Low back pain",
    "R07.9":    "Chest pain, unspecified",
    "K92.1":    "Gastrointestinal hemorrhage",
    "S83.511A": "Sprain of anterior cruciate ligament of right knee",
}


# ── File helpers ──────────────────────────────────────────────────────────────

def _allowed_file(filename: str) -> bool:
    ext = Path(filename).suffix.lower().lstrip(".")
    return ext in settings.ALLOWED_EXTENSIONS


async def _save_upload(file: UploadFile, subfolder: str) -> tuple[str, str]:
    """Save uploaded file with UUID prefix. Returns (original_name, abs_path)."""
    original  = file.filename or "upload"
    safe      = re.sub(r"[^\w.\-]", "_", original)
    unique    = f"{uuid.uuid4().hex}_{safe}"
    dest_dir  = settings.upload_dir / subfolder
    dest_dir.mkdir(parents=True, exist_ok=True)
    dest_path = dest_dir / unique

    contents = await file.read()
    dest_path.write_bytes(contents)
    file.file = io.BytesIO(contents)

    return original, str(dest_path)


def _extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from PDF / TXT / DOCX. Returns empty string on failure."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                return "\n".join((page.extract_text() or "") for page in pdf.pages)
        except Exception as exc:
            print(f"[PDF] Extraction failed — {exc}")
            return ""

    if ext == ".txt":
        return file_bytes.decode("utf-8", errors="replace")

    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as exc:
            print(f"[DOCX] Extraction failed — {exc}")
            return ""

    return ""


# ── Structured label-value parser ─────────────────────────────────────────────

# All labels that appear in our generated prescription PDF, in order.
# Each label maps to a field key. Values are separated by whitespace only
# (no colons) — the PDF renderer places label and value on the same line.
_PDF_LABELS = [
    ("Prescription ID",             "prescription_id"),
    ("Insurer",                     "insurer"),
    ("Source",                      "source"),
    ("Created At",                  "created_at"),
    ("Full Name",                   "patient_name"),
    ("Age",                         "age"),
    ("Gender",                      "gender"),
    ("Patient ID",                  "patient_id"),
    ("Insurance ID",                "insurance_id"),
    ("CPT Code",                    "cpt_code"),
    ("Procedure",                   "procedure_name"),
    ("Category",                    "procedure_category"),
    ("ICD-10 Code",                 "icd10_code"),
    ("Diagnosis",                   "diagnosis_desc"),
    ("BMI",                         "bmi"),
    ("Conservative Therapy (wks)",  "therapy_weeks"),
    ("Imaging Completed",           "imaging_completed"),
    ("Prerequisites Met",           "prerequisites_met"),
    ("Physician",                   "physician_name"),
    ("Specialty",                   "physician_specialty"),
    ("License No.",                 "physician_license"),
]

# Sort longest labels first so "Conservative Therapy (wks)" is matched
# before a hypothetical shorter prefix like "Conservative".
_PDF_LABELS_SORTED = sorted(_PDF_LABELS, key=lambda x: len(x[0]), reverse=True)


def _parse_label_value(text: str) -> dict:
    """
    Parse a prescription PDF using exact label matching.

    The prescription PDF produced by generate_prescription_pdf_bytes() puts
    each label and its value on the same line, separated only by whitespace:
        Full Name Noura Al-Zahrani
        CPT Code 73721

    This parser tries each known label against every line and extracts the
    remainder as the value. It is completely immune to the problems that
    plagued the old regex approach:
      - MD-69079 (license) no longer triggers the CPT 5-digit regex.
      - Names without colons are extracted correctly.
      - All clinical fields (therapy weeks, imaging, prerequisites) are parsed.
    """
    result = {}
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        for label, key in _PDF_LABELS_SORTED:
            if line.startswith(label):
                value = line[len(label):].strip()
                if value and key not in result:   # first match wins
                    result[key] = value
                break
    return result


def _parse_prescription_from_text(text: str, insurer: str) -> dict:
    """
    Parse every clinical field from an extracted prescription document.

    Strategy (in priority order):
      1. Structured label-value parse — exact match against our PDF labels.
         Handles every prescription generated by this system with 100% accuracy.
      2. Fallback regex scan — for third-party / hand-written documents that
         don't follow the label-value layout.

    Returns a flat dict with all fields needed to build a Prescription record.
    """
    # ── Step 1: structured parse ──────────────────────────────────────────────
    lv = _parse_label_value(text)

    # Patient name — direct from label parse
    patient_name = lv.get("patient_name", "").strip() or _fallback_name(text)

    # Age
    patient_age = 30
    if lv.get("age"):
        try:
            patient_age = min(int(lv["age"]), 120)
        except ValueError:
            patient_age = _fallback_age(text)
    else:
        patient_age = _fallback_age(text)

    # Gender
    gender_raw = lv.get("gender", "").strip()
    if gender_raw.lower() == "female":
        patient_gender = "Female"
    elif gender_raw.lower() == "male":
        patient_gender = "Male"
    else:
        patient_gender = _fallback_gender(text)

    # Patient IDs — from label parse
    patient_id   = lv.get("patient_id") or None
    insurance_id = lv.get("insurance_id") or None

    # CPT code — from label parse (avoids picking up MD-XXXXX license numbers)
    cpt_code = lv.get("cpt_code", "").strip()
    if not cpt_code:
        cpt_code = _fallback_cpt(text, exclude_after="License No")

    # Procedure name and category — from label parse, then catalogue lookup
    procedure_name     = lv.get("procedure_name", "").strip()
    procedure_category = lv.get("procedure_category", "").strip()
    if cpt_code and cpt_code in _PROCEDURE_META:
        meta = _PROCEDURE_META[cpt_code]
        procedure_name     = procedure_name     or meta["name"]
        procedure_category = procedure_category or meta["category"]
    procedure_name     = procedure_name     or "Extracted Procedure"
    procedure_category = procedure_category or "General"

    # ICD-10 code — from label parse
    icd10_code = lv.get("icd10_code", "").strip()
    if not icd10_code:
        icd10_code = _fallback_icd10(text)

    # Diagnosis — from label parse, then catalogue lookup
    diagnosis_desc = lv.get("diagnosis_desc", "").strip()
    if not diagnosis_desc and icd10_code in _DIAGNOSIS_META:
        diagnosis_desc = _DIAGNOSIS_META[icd10_code]
    diagnosis_desc = diagnosis_desc or "Diagnosis from uploaded document"

    # BMI
    bmi = None
    if lv.get("bmi"):
        try:
            bmi = float(lv["bmi"])
        except ValueError:
            bmi = _fallback_bmi(text)
    else:
        bmi = _fallback_bmi(text)

    # Conservative therapy weeks
    therapy_weeks = None
    if lv.get("therapy_weeks"):
        try:
            therapy_weeks = int(lv["therapy_weeks"])
        except ValueError:
            pass

    # Imaging completed
    imaging_raw = lv.get("imaging_completed", "").strip().lower()
    imaging_completed = imaging_raw in ("yes", "true", "1", "completed")

    # Prerequisites met — may be a semicolon-separated list on one line
    prereqs_raw = lv.get("prerequisites_met", "").strip()
    if prereqs_raw and prereqs_raw.lower() not in ("none", "none documented", "—"):
        # Split on semicolon if multiple prerequisites are listed
        prerequisites_met = [p.strip() for p in prereqs_raw.split(";") if p.strip()]
    else:
        prerequisites_met = []

    # Physician fields
    physician_name      = lv.get("physician_name", "").strip() or None
    physician_specialty = lv.get("physician_specialty", "").strip() or None
    physician_license   = lv.get("physician_license", "").strip() or None

    return {
        "patient_name":            patient_name,
        "patient_id":              patient_id,
        "patient_age":             patient_age,
        "patient_gender":          patient_gender,
        "insurance_id":            insurance_id,
        "insurer":                 lv.get("insurer", insurer) or insurer,
        "cpt_code":                cpt_code or "99999",
        "procedure_name":          procedure_name,
        "procedure_category":      procedure_category,
        "icd10_code":              icd10_code or "Z00.00",
        "diagnosis_desc":          diagnosis_desc,
        "bmi":                     bmi,
        "conservative_therapy_weeks": therapy_weeks,
        "imaging_completed":       imaging_completed,
        "prerequisites_met":       prerequisites_met,
        "physician_name":          physician_name,
        "physician_specialty":     physician_specialty,
        "physician_license":       physician_license,
    }


# ── Regex fallbacks (for non-standard / third-party documents) ────────────────

def _fallback_name(text: str) -> str:
    """Extract patient name using flexible patterns for non-standard documents."""
    patterns = [
        r"(?:Full\s+Name|Patient\s+Name|Name)[:\s]+([A-Z][a-z]+(?:\s+[A-Z][a-zA-Z\-']+){1,4})",
        r"(?:Patient|Name)[:\s]+([A-Z][a-z]+\s+[A-Z][a-zA-Z]+)",
    ]
    for pat in patterns:
        m = re.search(pat, text)
        if m:
            return m.group(1).strip()
    return "Unknown Patient"


def _fallback_age(text: str) -> int:
    m = re.search(r"\bAge[:\s]+(\d{1,3})\b", text, re.IGNORECASE)
    if m:
        try:
            return min(int(m.group(1)), 120)
        except ValueError:
            pass
    return 30


def _fallback_gender(text: str) -> str:
    if re.search(r"\bfemale\b", text, re.IGNORECASE):
        return "Female"
    if re.search(r"\bmale\b", text, re.IGNORECASE):
        return "Male"
    return "Unknown"


def _fallback_bmi(text: str) -> Optional[float]:
    m = re.search(r"\bBMI[:\s]+([\d.]+)\b", text, re.IGNORECASE)
    if m:
        try:
            return float(m.group(1))
        except ValueError:
            pass
    return None


def _fallback_cpt(text: str, exclude_after: str = "") -> str:
    """
    Find the first 5-digit CPT code in the text, excluding anything that
    appears after a known non-CPT section marker (e.g. 'License No.').
    Prevents MD-XXXXX license numbers from being picked up as CPT codes.
    """
    # Truncate text at the exclusion marker if provided
    if exclude_after:
        idx = text.find(exclude_after)
        if idx > 0:
            text = text[:idx]

    # Look for explicit 'CPT Code XXXXX' pattern first
    m = re.search(r"\bCPT\s+(?:Code\s+)?(\d{5})\b", text, re.IGNORECASE)
    if m:
        return m.group(1)

    # Fall back to any standalone 5-digit number
    m = re.search(r"\b(\d{5})\b", text)
    return m.group(1) if m else "99999"


def _fallback_icd10(text: str) -> str:
    ICD10_PATTERN = r"\b[A-Z]\d{2}(?:\.\d{1,4})?\b"
    codes = re.findall(ICD10_PATTERN, text)
    return codes[0] if codes else "Z00.00"


# ── URL helpers ───────────────────────────────────────────────────────────────

def build_download_url(base_url: str, prescription_db_id: int) -> str:
    return f"{base_url.rstrip('/')}/api/download/prescription/{prescription_db_id}"


def build_prescription_dict(record: Prescription, base_url: str) -> dict:
    d = record.to_dict()
    d["download_url"] = build_download_url(base_url, record.id)
    return d


# ── PDF generation ────────────────────────────────────────────────────────────

def generate_prescription_pdf_bytes(record: Prescription) -> bytes:
    """Build a formatted PDF for a prescription record."""
    try:
        from reportlab.lib.pagesizes  import A4
        from reportlab.lib.units      import cm
        from reportlab.lib            import colors
        from reportlab.platypus       import (
            SimpleDocTemplate, Paragraph, Spacer, Table, HRFlowable,
        )
        from reportlab.platypus.tables import TableStyle as TS
        from reportlab.lib.styles     import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums      import TA_CENTER
    except ImportError:
        raise ImportError("reportlab")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
    )

    base  = getSampleStyleSheet()
    BRAND = colors.HexColor("#1a56db")

    st = {
        "title":    ParagraphStyle("title",    parent=base["Title"],
                                   textColor=BRAND, fontSize=18, spaceAfter=4),
        "subtitle": ParagraphStyle("subtitle", parent=base["Normal"],
                                   textColor=colors.HexColor("#4b5563"),
                                   fontSize=10, spaceAfter=12),
        "heading":  ParagraphStyle("heading",  parent=base["Heading2"],
                                   textColor=BRAND, fontSize=11, spaceBefore=10, spaceAfter=4),
        "body":     ParagraphStyle("body",     parent=base["Normal"], fontSize=9, leading=13),
        "small":    ParagraphStyle("small",    parent=base["Normal"],
                                   fontSize=7.5, textColor=colors.HexColor("#6b7280")),
    }

    def _ts():
        return TS([
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f3f4f6")),
            ("TEXTCOLOR",  (0, 0), (0, -1), colors.HexColor("#374151")),
            ("FONTSIZE",   (0, 0), (-1, -1), 9),
            ("GRID",       (0, 0), (-1, -1), 0.4, colors.HexColor("#e5e7eb")),
            ("VALIGN",     (0, 0), (-1, -1), "TOP"),
            ("PADDING",    (0, 0), (-1, -1), 5),
        ])

    story = []

    # Header
    story.append(Paragraph("PA Validation System", st["subtitle"]))
    story.append(Paragraph("Patient Prescription", st["title"]))
    story.append(HRFlowable(width="100%", thickness=1.5, color=BRAND))
    story.append(Spacer(1, 0.3*cm))

    meta_rows = [
        ["Prescription ID", record.prescription_id],
        ["Insurer",         record.insurer],
        ["Source",          record.source.replace("_", " ").title()],
        ["Created At",      record.created_at.strftime("%Y-%m-%d %H:%M UTC")],
    ]
    if record.filename:
        meta_rows.append(["Uploaded File", record.filename])

    mt = Table(meta_rows, colWidths=[5*cm, 11*cm])
    mt.setStyle(_ts())
    story.append(mt)
    story.append(Spacer(1, 0.4*cm))

    # Patient
    story.append(Paragraph("Patient Information", st["heading"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#aaaaaa")))
    story.append(Spacer(1, 0.1*cm))
    patient_rows = [
        ["Full Name",    record.patient_name],
        ["Age",          str(record.patient_age)],
        ["Gender",       record.patient_gender],
        ["Patient ID",   record.patient_id or "—"],
        ["Insurance ID", record.insurance_id or "—"],
    ]
    pt = Table(patient_rows, colWidths=[5*cm, 11*cm])
    pt.setStyle(_ts())
    story.append(pt)
    story.append(Spacer(1, 0.3*cm))

    # Procedure & Diagnosis
    story.append(Paragraph("Prescribed Procedure & Diagnosis", st["heading"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#aaaaaa")))
    story.append(Spacer(1, 0.1*cm))
    proc_rows = [
        ["CPT Code",    record.cpt_code],
        ["Procedure",   record.procedure_name],
        ["Category",    record.procedure_category],
        ["ICD-10 Code", record.icd10_code],
        ["Diagnosis",   record.diagnosis_desc],
    ]
    prt = Table(proc_rows, colWidths=[5*cm, 11*cm])
    prt.setStyle(_ts())
    story.append(prt)
    story.append(Spacer(1, 0.3*cm))

    # Clinical info
    story.append(Paragraph("Clinical Information", st["heading"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#aaaaaa")))
    story.append(Spacer(1, 0.1*cm))
    prereqs     = record.prerequisites_met or []
    prereq_str  = "; ".join(prereqs) if prereqs else "None documented"
    clin_rows = [
        ["BMI",                        str(record.bmi) if record.bmi else "—"],
        ["Conservative Therapy (wks)", str(record.conservative_therapy_weeks)
                                       if record.conservative_therapy_weeks is not None else "—"],
        ["Imaging Completed",          "Yes" if record.imaging_completed else "No"],
        ["Prerequisites Met",          prereq_str],
    ]
    ct = Table(clin_rows, colWidths=[5*cm, 11*cm])
    ct.setStyle(_ts())
    story.append(ct)
    story.append(Spacer(1, 0.3*cm))

    # Physician
    if record.physician_name:
        story.append(Paragraph("Requesting Physician", st["heading"]))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#aaaaaa")))
        story.append(Spacer(1, 0.1*cm))
        phys_rows = [
            ["Physician",   record.physician_name],
            ["Specialty",   record.physician_specialty or "—"],
            ["License No.", record.physician_license or "—"],
        ]
        phyt = Table(phys_rows, colWidths=[5*cm, 11*cm])
        phyt.setStyle(_ts())
        story.append(phyt)
        story.append(Spacer(1, 0.3*cm))

    # Footer
    story.append(Spacer(1, 0.5*cm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#aaaaaa")))
    story.append(Spacer(1, 0.2*cm))
    story.append(Paragraph(
        "This document was generated by the PA Validation System for testing and "
        "demonstration purposes only. It does not constitute a legally binding medical "
        "prescription.  PA Validation System | Queen Mary University of London | 2026",
        st["small"],
    ))

    doc.build(story)
    buf.seek(0)
    return buf.read()


# ── Service functions ─────────────────────────────────────────────────────────

async def generate_prescription(
    db:         AsyncSession,
    insurer:    Optional[str],
    likelihood: str,
    base_url:   str,
) -> dict:
    """
    Generate a synthetic PA prescription from the data generator,
    persisting every field from the synthetic_pa_requests.json schema.
    """
    from src.nlp.data_generator import INSURERS, generate_policy as _gen_policy
    import random

    if not insurer:
        insurer = random.choice(INSURERS)

    temp_policy = _gen_policy(insurer=insurer, num_procedures=5)

    pa = generator.generate_pa_request(
        policy=temp_policy,
        approval_likelihood=likelihood if likelihood != "random" else "random",
    )

    patient = pa.get("patient", {})
    proc    = pa.get("requested_procedure", {})
    diag    = pa.get("diagnosis", {})
    clin    = pa.get("clinical_info", {})
    phys    = pa.get("requesting_physician", {})

    rx_id = f"RX-{uuid.uuid4().hex[:6].upper()}"

    record = Prescription(
        prescription_id            = rx_id,
        source                     = "generated",
        filename                   = None,
        file_path                  = None,
        patient_id                 = patient.get("patient_id"),
        patient_name               = patient.get("name", "Unknown"),
        patient_age                = patient.get("age", 30),
        patient_gender             = patient.get("gender", "Unknown"),
        insurance_id               = patient.get("insurance_id"),
        insurer                    = pa.get("insurer", insurer),
        cpt_code                   = proc.get("cpt_code", "99999"),
        procedure_name             = proc.get("name", "Unknown Procedure"),
        procedure_category         = proc.get("category", "General"),
        icd10_code                 = diag.get("icd10_code", "Z00.00"),
        diagnosis_desc             = diag.get("description", "Unknown"),
        bmi                        = clin.get("bmi"),
        conservative_therapy_weeks = clin.get("conservative_therapy_duration_weeks"),
        imaging_completed          = clin.get("imaging_completed", False),
        prerequisites_met          = clin.get("prerequisites_met", []),
        physician_name             = phys.get("name"),
        physician_specialty        = phys.get("specialty"),
        physician_license          = phys.get("license_number"),
        approval_likelihood_label  = pa.get("approval_likelihood_label"),
        full_data                  = pa,
    )
    db.add(record)
    await db.commit()
    await db.refresh(record)

    return build_prescription_dict(record, base_url)


async def upload_prescription(
    db:       AsyncSession,
    file:     UploadFile,
    insurer:  str,
    base_url: str,
) -> dict:
    """
    Upload a prescription document (PDF / DOCX / TXT), extract every clinical
    field using the structured label-value parser, and persist the record.
    """
    if not _allowed_file(file.filename or ""):
        raise HTTPException(
            status_code=400,
            detail=f"File type not allowed. Accepted: {', '.join(sorted(settings.ALLOWED_EXTENSIONS))}.",
        )

    file_bytes    = await file.read()
    original_name = file.filename or "upload"

    extracted_text = _extract_text(file_bytes, original_name)
    if not extracted_text.strip():
        raise HTTPException(
            status_code=422,
            detail=(
                f"No readable text could be extracted from '{original_name}'. "
                "Please upload a text-selectable PDF, DOCX, or TXT file."
            ),
        )

    parsed = _parse_prescription_from_text(extracted_text, insurer)

    print(
        f"[rx-upload] '{original_name}' — "
        f"patient={parsed['patient_name']!r}, "
        f"CPT={parsed['cpt_code']}, ICD10={parsed['icd10_code']}, "
        f"BMI={parsed['bmi']}, therapy_wks={parsed['conservative_therapy_weeks']}, "
        f"imaging={parsed['imaging_completed']}"
    )

    # Save file to disk after successful parse
    file.file = io.BytesIO(file_bytes)
    original_name, saved_path = await _save_upload(file, subfolder="prescriptions")

    rx_id = f"RX-{uuid.uuid4().hex[:6].upper()}"

    record = Prescription(
        prescription_id            = rx_id,
        source                     = "uploaded",
        filename                   = original_name,
        file_path                  = saved_path,
        patient_id                 = parsed.get("patient_id"),
        patient_name               = parsed["patient_name"],
        patient_age                = parsed["patient_age"],
        patient_gender             = parsed["patient_gender"],
        insurance_id               = parsed.get("insurance_id"),
        insurer                    = parsed["insurer"],
        cpt_code                   = parsed["cpt_code"],
        procedure_name             = parsed["procedure_name"],
        procedure_category         = parsed["procedure_category"],
        icd10_code                 = parsed["icd10_code"],
        diagnosis_desc             = parsed["diagnosis_desc"],
        bmi                        = parsed["bmi"],
        conservative_therapy_weeks = parsed["conservative_therapy_weeks"],
        imaging_completed          = parsed["imaging_completed"],
        prerequisites_met          = parsed["prerequisites_met"],
        physician_name             = parsed["physician_name"],
        physician_specialty        = parsed["physician_specialty"],
        physician_license          = parsed["physician_license"],
        approval_likelihood_label  = None,
        full_data                  = parsed,
    )
    db.add(record)
    await db.commit()
    await db.refresh(record)

    return build_prescription_dict(record, base_url)


async def list_prescriptions(db: AsyncSession, base_url: str) -> list:
    """Return all prescriptions, newest first."""
    result  = await db.execute(
        select(Prescription).order_by(Prescription.created_at.desc())
    )
    records = result.scalars().all()
    return [build_prescription_dict(r, base_url) for r in records]


async def get_prescription_or_404(db: AsyncSession, prescription_db_id: int) -> Prescription:
    """Fetch Prescription by PK or raise HTTP 404."""
    result = await db.execute(
        select(Prescription).where(Prescription.id == prescription_db_id)
    )
    record = result.scalar_one_or_none()
    if record is None:
        raise HTTPException(status_code=404, detail="Prescription not found.")
    return record